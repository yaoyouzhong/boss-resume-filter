from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from legacy_doc_converter import LegacyDocConversionError
from resume_parser import (
    ResumeContentTooShortError,
    ResumeParserDependencyError,
    ResumeTextReadError,
    UnsupportedResumeFormatError,
    parse_resume_text,
)


def _capture_exception(exception_type, callback):
    try:
        callback()
    except exception_type as exc:
        return exc
    raise AssertionError(f"未抛出 {exception_type.__name__}")


def test_parse_resume_text_reads_utf8_and_strips_outer_whitespace():
    content = "  " + "Python 后端开发经验，负责交易系统、数据库和自动化交付。" * 3 + "  "
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.txt"
        path.write_text(content, encoding="utf-8")

        parsed = parse_resume_text(path)

    assert parsed == content.strip()


def test_parse_resume_text_reads_docx_table_cells_and_textboxes():
    """表格排版简历的正文在表格单元格里，document.paragraphs 读不到。"""
    import docx

    document = docx.Document()
    document.add_paragraph("卞婷婷 求职简历")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "卞婷婷"
    table.cell(1, 0).text = "工作经验"
    table.cell(1, 1).text = "6 年 Java 开发，熟悉 Spring、MySQL 与银行核心系统"
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.docx"
        document.save(str(path))

        parsed = parse_resume_text(path, minimum_length=10)

    assert "卞婷婷 求职简历" in parsed
    assert "6 年 Java 开发" in parsed  # 表格单元格内容必须出现
    assert "Spring" in parsed


def test_parse_resume_text_falls_back_to_gbk_for_markdown():
    content = "Java 开发工程师，熟悉金融交易、数据库设计、故障排查和项目交付。" * 3
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.md"
        path.write_bytes(content.encode("gbk"))

        parsed = parse_resume_text(path)

    assert parsed == content


def test_parse_resume_text_removes_html_code_and_restores_entities():
    visible = "候选人 & 简历：十年软件开发经验，熟悉 Python、SQL 和证券交易系统。" * 3
    document = (
        "<html><style>.hidden {display:none}</style>"
        "<script>secret();</script><body>"
        f"<p>{visible.replace('&', '&amp;')}</p></body></html>"
    )
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.html"
        path.write_text(document, encoding="utf-8")

        parsed = parse_resume_text(path)

    assert parsed == visible
    assert "secret" not in parsed
    assert "hidden" not in parsed


def test_parse_resume_text_classifies_unsupported_format():
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.pages"
        path.write_text("x" * 100, encoding="utf-8")

        exc = _capture_exception(
            UnsupportedResumeFormatError,
            lambda: parse_resume_text(path),
        )

    assert exc.extension == ".pages"


def test_parse_resume_text_classifies_missing_optional_parser_dependency():
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.pdf"
        path.write_bytes(b"not-read-when-parser-is-missing")
        with patch.dict("sys.modules", {"pdfminer.high_level": None}):
            exc = _capture_exception(
                ResumeParserDependencyError,
                lambda: parse_resume_text(path),
            )

    assert exc.format_name == "PDF"
    assert exc.package_name == "pdfminer.six"


def test_parse_resume_text_classifies_empty_text_file():
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.txt"
        path.write_text("", encoding="utf-8")

        exc = _capture_exception(
            ResumeTextReadError,
            lambda: parse_resume_text(path),
        )

    assert exc.format_name == "文本"


def test_parse_resume_text_reports_short_extracted_content():
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.txt"
        path.write_text("简历内容过少", encoding="utf-8")

        exc = _capture_exception(
            ResumeContentTooShortError,
            lambda: parse_resume_text(path),
        )

    assert exc.text_length == 6
    assert exc.minimum_length == 50


def test_parse_resume_text_reads_legacy_doc_via_word_text_and_cleans_up():
    """真 OLE2 .doc 经本机 Word 提取文本，Word 换行/表格符归一化并清理临时目录。"""
    holder = {}
    raw_lines = [
        "李四\r5 年 Java 开发经验\r技能\tSpring\x07MySQL\r",
        "熟悉银行核心交易系统与批量处理。" * 3,
    ]

    def fake_convert(source_path):
        converted_dir = Path(source_path).parent / "converted-temp"
        converted_dir.mkdir()
        converted = converted_dir / "resume.txt"
        converted.write_text("".join(raw_lines), encoding="utf-8")
        holder["dir"] = converted_dir
        return converted

    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(b"\xd0\xcf\x11\xe0 legacy binary")
        with patch("resume_parser.convert_legacy_doc", fake_convert):
            parsed = parse_resume_text(path)

    assert "\r" not in parsed
    assert "\x07" not in parsed
    assert "李四\n5 年 Java 开发经验" in parsed
    assert "Spring\tMySQL" in parsed
    assert not holder["dir"].exists()  # 临时目录已清理


def test_parse_resume_text_parses_html_disguised_as_doc_without_word():
    """HTML 伪装 .doc（招聘网站导出常见）必须直接解析，不启动本机 Word。"""
    html = (
        "<html xmlns:o=\"urn:schemas-microsoft-com:office:office\"><head>"
        "<style>.x{color:red}</style></head><body>"
        "<p>王必强 8 年后端开发，熟悉 Java、Spring、MySQL，"
        "参与银行核心与支付清算系统建设，主导过多个批量与联机交易模块。</p></body></html>"
    )

    def fail_convert(_source_path):
        raise AssertionError("伪格式 .doc 不应调用 Word 转换")

    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(html.encode("gbk"))
        with patch("resume_parser.convert_legacy_doc", fail_convert):
            parsed = parse_resume_text(path)

    assert "王必强" in parsed
    assert "Spring" in parsed
    assert "<p>" not in parsed


def test_parse_resume_text_parses_utf16_html_disguised_as_doc_without_word():
    """UTF-16 HTML 伪装 .doc 也必须直接解析，不能误送 Word COM。"""
    html = (
        "<html><body><p>周敏 7 年数据开发经验，熟悉 Python、SQL、ETL，"
        "参与银行数据仓库、监管报送和批量调度平台建设。</p></body></html>"
    )

    def fail_convert(_source_path):
        raise AssertionError("UTF-16 HTML 伪格式 .doc 不应调用 Word 转换")

    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(html.encode("utf-16"))
        with patch("resume_parser.convert_legacy_doc", fail_convert):
            parsed = parse_resume_text(path, minimum_length=10)

    assert "周敏" in parsed
    assert "数据仓库" in parsed


def test_parse_resume_text_parses_mht_disguised_as_doc_without_word():
    """MHT（MIME HTML）伪装 .doc 必须经 MIME 提取 HTML 部分解析。

    part 头不声明 charset（真实招聘导出版本如此），编码取自 HTML meta。
    """
    import quopri

    body = (
        "<html><head><meta http-equiv=Content-Type content=\"text/html; charset=utf-8\">"
        "</head><body><p>许敏 5 年需求分析经验，熟悉核心系统改造与数据迁移，"
        "参与信贷、支付、渠道整合等多个银行项目的需求梳理与落地。</p></body></html>"
    )
    encoded = quopri.encodestring(body.encode("utf-8")).decode("ascii")
    mht = (
        "MIME-Version: 1.0\r\n"
        "Content-Type: multipart/related; boundary=\"----=_NextPart_000\"\r\n"
        "\r\n"
        "------=_NextPart_000\r\n"
        "Content-Type: text/html\r\n"
        "Content-Transfer-Encoding: quoted-printable\r\n"
        "\r\n"
        + encoded.replace("\n", "\r\n")
        + "\r\n------=_NextPart_000--\r\n"
    )

    def fail_convert(_source_path):
        raise AssertionError("伪格式 .doc 不应调用 Word 转换")

    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(mht.encode("ascii"))
        with patch("resume_parser.convert_legacy_doc", fail_convert):
            parsed = parse_resume_text(path)

    assert "许敏" in parsed
    assert "需求分析" in parsed


def test_parse_resume_text_honors_mht_part_charset_before_html_meta():
    """MIME part 已声明字符集时应优先使用，避免无 meta 的中文变乱码。"""
    import quopri

    body = (
        "<html><body><p>赵强 9 年系统架构经验，熟悉 Java、分布式系统，"
        "负责证券交易与清算平台的架构设计和性能治理。</p></body></html>"
    )
    encoded = quopri.encodestring(body.encode("gbk")).decode("ascii")
    mht = (
        "MIME-Version: 1.0\r\n"
        "Content-Type: multipart/related; boundary=\"part-boundary\"\r\n\r\n"
        "--part-boundary\r\n"
        "Content-Type: text/html; charset=gbk\r\n"
        "Content-Transfer-Encoding: quoted-printable\r\n\r\n"
        + encoded.replace("\n", "\r\n")
        + "\r\n--part-boundary--\r\n"
    )
    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(mht.encode("ascii"))
        with patch(
            "resume_parser.convert_legacy_doc",
            side_effect=AssertionError("MHT 不应调用 Word 转换"),
        ):
            parsed = parse_resume_text(path, minimum_length=10)

    assert "赵强" in parsed
    assert "证券交易" in parsed


def test_parse_resume_text_propagates_legacy_doc_conversion_failure():
    """本机没有 Word 等转换失败必须原样向调用方抛出分类错误。"""
    def fake_convert(_source_path):
        raise LegacyDocConversionError("未检测到本机安装的 Microsoft Word")

    with TemporaryDirectory() as directory:
        path = Path(directory) / "resume.doc"
        path.write_bytes(b"\xd0\xcf\x11\xe0 legacy binary")
        with patch("resume_parser.convert_legacy_doc", fake_convert):
            exc = _capture_exception(
                LegacyDocConversionError,
                lambda: parse_resume_text(path),
            )

    assert exc.reason == "未检测到本机安装的 Microsoft Word"
